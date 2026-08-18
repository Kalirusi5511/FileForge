#!/usr/bin/env python3
"""
AI File Platform - Projekt-Erstellungs-Skript
Dieses Skript erstellt alle notwendigen Dateien und Ordner für die Anwendung.
"""

import os
import sys
import shutil
from pathlib import Path

# Projekt-Struktur
PROJECT_STRUCTURE = {
    '': [
        'app.py',
        'requirements.txt',
        'README.md'
    ],
    'converters': [
        '__init__.py',
        'archive.py',
        'image.py',
        'office.py',
        'pdf.py',
        'text.py',
        'router.py'
    ],
    'templates': [
        'index.html'
    ]
}

def create_project():
    """Erstellt das gesamte Projekt"""
    print("🚀 Erstelle AI File Platform Projekt...")
    
    # Erstelle Projektordner
    project_root = Path.cwd() / 'ai-file-platform'
    project_root.mkdir(exist_ok=True)
    
    # Wechsle in den Projektordner
    os.chdir(project_root)
    print(f"📁 Projektverzeichnis: {project_root}")
    
    # Erstelle Ordner und Dateien
    for folder, files in PROJECT_STRUCTURE.items():
        if folder:
            folder_path = Path(folder)
            folder_path.mkdir(exist_ok=True)
            print(f"  📁 Erstelle Ordner: {folder}")
        
        for filename in files:
            file_path = Path(folder) / filename if folder else Path(filename)
            
            # Generiere Inhalt basierend auf Dateiname
            content = get_file_content(folder, filename)
            
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            print(f"  📄 Erstelle Datei: {file_path}")
    
    # Erstelle leere Ordner
    empty_folders = ['uploads', 'outputs', 'static']
    for folder in empty_folders:
        Path(folder).mkdir(exist_ok=True)
        print(f"  📁 Erstelle Ordner: {folder}")
    
    print("\n✅ Projekt erfolgreich erstellt!")
    print(f"\n📂 Projektverzeichnis: {project_root}")
    print("\n🚀 Nächste Schritte:")
    print("  1. cd ai-file-platform")
    print("  2. python -m venv venv")
    print("  3. source venv/bin/activate  # oder venv\\Scripts\\activate auf Windows")
    print("  4. pip install -r requirements.txt")
    print("  5. python app.py")
    print("  6. Öffne http://localhost:5001 im Browser")

def get_file_content(folder, filename):
    """Gibt den Inhalt für eine bestimmte Datei zurück"""
    
    # app.py
    if filename == 'app.py':
        return '''import os
import zipfile
import shutil
import uuid
import threading
import time
from flask import Flask, render_template, request, send_file, jsonify, after_this_request
from werkzeug.utils import secure_filename
from converters.router import ConverterRouter
from pathlib import Path

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'outputs'
app.config['MAX_CONTENT_LENGTH'] = 500 * 1024 * 1024  # 500MB max
app.config['SECRET_KEY'] = os.urandom(24)

# Erstelle Ordner
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)
os.makedirs('static', exist_ok=True)

# Globale Session-Status-Speicher
sessions = {}

class ConversionSession:
    def __init__(self, session_id):
        self.session_id = session_id
        self.upload_dir = os.path.join(app.config['UPLOAD_FOLDER'], session_id)
        self.output_dir = os.path.join(app.config['OUTPUT_FOLDER'], session_id)
        self.files = []
        self.converted_files = []
        self.status = 'processing'
        self.progress = 0
        self.total_files = 0
        self.processed_files = 0
        self.errors = []
        
        os.makedirs(self.upload_dir, exist_ok=True)
        os.makedirs(self.output_dir, exist_ok=True)
    
    def add_file(self, filename, filepath):
        self.files.append({
            'name': filename,
            'path': filepath,
            'status': 'pending'
        })
        self.total_files += 1
    
    def update_progress(self):
        self.progress = int((self.processed_files / self.total_files) * 100) if self.total_files > 0 else 0
    
    def get_status(self):
        return {
            'status': self.status,
            'progress': self.progress,
            'total_files': self.total_files,
            'processed_files': self.processed_files,
            'files': self.files,
            'converted_files': self.converted_files,
            'errors': self.errors
        }

def is_ignored(filename):
    """Überprüft ob eine Datei ignoriert werden soll"""
    ignored_patterns = [
        '.pyc', '.dll', '.exe', '.so',
        '__pycache__', '.git', '.venv', 'venv',
        '.DS_Store', 'Thumbs.db', '.idea', '.vscode',
        'node_modules', 'vendor', 'dist', 'build'
    ]
    
    for pattern in ignored_patterns:
        if pattern in filename:
            return True
    return False

def process_zip_file(zip_path, session):
    """Extrahiert eine ZIP-Datei und fügt die enthaltenen Dateien zur Session hinzu"""
    try:
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            extract_dir = os.path.join(session.upload_dir, 'extracted')
            os.makedirs(extract_dir, exist_ok=True)
            zip_ref.extractall(extract_dir)
            
            # Verarbeite alle extrahierten Dateien
            for root, dirs, files in os.walk(extract_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    rel_path = os.path.relpath(file_path, extract_dir)
                    
                    if not is_ignored(rel_path):
                        # Kopiere Datei in den Upload-Ordner
                        dest_path = os.path.join(session.upload_dir, rel_path)
                        os.makedirs(os.path.dirname(dest_path), exist_ok=True)
                        shutil.copy2(file_path, dest_path)
                        session.add_file(rel_path, dest_path)
            
            # Lösche temporären Extraktionsordner
            shutil.rmtree(extract_dir)
            return True
            
    except Exception as e:
        session.errors.append(f"Fehler beim Entpacken von {zip_path}: {str(e)}")
        return False

def process_session(session_id):
    """Verarbeitet alle Dateien in einer Session"""
    session = sessions.get(session_id)
    if not session:
        return
    
    router = ConverterRouter()
    
    try:
        # Verarbeite alle Dateien
        for file_info in session.files:
            try:
                file_path = file_info['path']
                file_name = file_info['name']
                
                # Prüfe ob Datei konvertiert werden kann
                if router.can_convert(file_name):
                    # Konvertiere Datei
                    output_filename = f"{os.path.splitext(file_name)[0]}.txt"
                    output_path = os.path.join(session.output_dir, output_filename)
                    
                    # Vermeide Überschreiben
                    counter = 1
                    while os.path.exists(output_path):
                        output_filename = f"{os.path.splitext(file_name)[0]}_{counter}.txt"
                        output_path = os.path.join(session.output_dir, output_filename)
                        counter += 1
                    
                    success = router.convert(file_path, output_path)
                    
                    if success and os.path.exists(output_path):
                        session.converted_files.append({
                            'original': file_name,
                            'converted': output_filename
                        })
                        file_info['status'] = 'success'
                    else:
                        file_info['status'] = 'error'
                        session.errors.append(f"Fehler beim Konvertieren von {file_name}")
                else:
                    file_info['status'] = 'skipped'
                    
            except Exception as e:
                file_info['status'] = 'error'
                session.errors.append(f"Fehler bei {file_info['name']}: {str(e)}")
            
            session.processed_files += 1
            session.update_progress()
            
        session.status = 'completed'
        
    except Exception as e:
        session.status = 'error'
        session.errors.append(f"Allgemeiner Fehler: {str(e)}")

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_files():
    """Verarbeitet Datei-Uploads"""
    session_id = str(uuid.uuid4())
    session = ConversionSession(session_id)
    sessions[session_id] = session
    
    files = request.files.getlist('files')
    
    for file in files:
        if file.filename == '':
            continue
            
        filename = secure_filename(file.filename)
        
        # Prüfe ob Datei ignoriert werden soll
        if is_ignored(filename):
            continue
            
        file_path = os.path.join(session.upload_dir, filename)
        file.save(file_path)
        
        # Prüfe ob es eine ZIP-Datei ist
        if filename.lower().endswith('.zip'):
            process_zip_file(file_path, session)
        else:
            session.add_file(filename, file_path)
    
    # Starte Verarbeitung in Hintergrund-Thread
    thread = threading.Thread(target=process_session, args=(session_id,))
    thread.daemon = True
    thread.start()
    
    return jsonify({'session_id': session_id})

@app.route('/status/<session_id>')
def get_status(session_id):
    """Gibt den aktuellen Verarbeitungsstatus zurück"""
    session = sessions.get(session_id)
    if not session:
        return jsonify({'error': 'Session nicht gefunden'}), 404
    
    return jsonify(session.get_status())

@app.route('/download/<session_id>')
def download_files(session_id):
    """Lädt alle konvertierten Dateien als ZIP herunter"""
    session = sessions.get(session_id)
    if not session:
        return jsonify({'error': 'Session nicht gefunden'}), 404
    
    if session.status != 'completed':
        return jsonify({'error': 'Konvertierung noch nicht abgeschlossen'}), 400
    
    if not session.converted_files:
        return jsonify({'error': 'Keine konvertierten Dateien verfügbar'}), 400
    
    # Erstelle ZIP-Datei
    zip_path = os.path.join(app.config['OUTPUT_FOLDER'], f"{session_id}.zip")
    
    with zipfile.ZipFile(zip_path, 'w') as zipf:
        for converted_file in session.converted_files:
            file_path = os.path.join(session.output_dir, converted_file['converted'])
            if os.path.exists(file_path):
                zipf.write(file_path, converted_file['converted'])
    
    # Lösche Session nach Download
    @after_this_request
    def cleanup(response):
        try:
            # Lösche temporäre Dateien
            shutil.rmtree(session.upload_dir, ignore_errors=True)
            shutil.rmtree(session.output_dir, ignore_errors=True)
            if os.path.exists(zip_path):
                os.remove(zip_path)
            if session_id in sessions:
                del sessions[session_id]
        except Exception as e:
            print(f"Cleanup Fehler: {e}")
        return response
    
    return send_file(zip_path, as_attachment=True, download_name=f"converted_files_{session_id}.zip")

@app.route('/download/<session_id>/<filename>')
def download_single_file(session_id, filename):
    """Lädt eine einzelne konvertierte Datei herunter"""
    session = sessions.get(session_id)
    if not session:
        return jsonify({'error': 'Session nicht gefunden'}), 404
    
    file_path = os.path.join(session.output_dir, filename)
    if not os.path.exists(file_path):
        return jsonify({'error': 'Datei nicht gefunden'}), 404
    
    return send_file(file_path, as_attachment=True)

@app.route('/preview/<session_id>/<filename>')
def preview_file(session_id, filename):
    """Zeigt eine Vorschau der konvertierten Datei"""
    session = sessions.get(session_id)
    if not session:
        return jsonify({'error': 'Session nicht gefunden'}), 404
    
    file_path = os.path.join(session.output_dir, filename)
    if not os.path.exists(file_path):
        return jsonify({'error': 'Datei nicht gefunden'}), 404
    
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
            # Begrenze Vorschau auf 5000 Zeichen
            if len(content) > 5000:
                content = content[:5000] + '\\n\\n... (Vorschau auf 5000 Zeichen begrenzt)'
            return jsonify({'content': content})
    except UnicodeDecodeError:
        return jsonify({'error': 'Datei kann nicht als Text angezeigt werden'}), 400
    except Exception as e:
        return jsonify({'error': f'Fehler beim Lesen der Datei: {str(e)}'}), 400

@app.route('/cleanup/<session_id>', methods=['POST'])
def cleanup_session(session_id):
    """Bereinigt eine Session manuell"""
    session = sessions.get(session_id)
    if session:
        try:
            shutil.rmtree(session.upload_dir, ignore_errors=True)
            shutil.rmtree(session.output_dir, ignore_errors=True)
            del sessions[session_id]
            return jsonify({'success': True})
        except Exception as e:
            return jsonify({'error': str(e)}), 400
    return jsonify({'error': 'Session nicht gefunden'}), 404

@app.errorhandler(413)
def too_large(e):
    return jsonify({'error': 'Datei zu groß. Maximale Größe: 500MB'}), 413

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5001)'''

    # converters/__init__.py
    if folder == 'converters' and filename == '__init__.py':
        return '''"""
Converter-Module für die AI File Platform

Dieses Modul enthält alle Konverter für verschiedene Dateiformate.
"""

from .pdf import PDFConverter
from .office import OfficeConverter
from .image import ImageConverter
from .text import TextConverter
from .archive import ArchiveConverter
from .router import ConverterRouter

__all__ = [
    'PDFConverter',
    'OfficeConverter', 
    'ImageConverter',
    'TextConverter',
    'ArchiveConverter',
    'ConverterRouter'
]'''

    # converters/pdf.py
    if folder == 'converters' and filename == 'pdf.py':
        return '''import os
import PyPDF2

class PDFConverter:
    """Konverter für PDF-Dateien"""
    
    def convert(self, input_path, output_path):
        """
        Konvertiert eine PDF-Datei zu TXT
        
        Args:
            input_path (str): Pfad zur PDF-Datei
            output_path (str): Pfad zur Ausgabedatei
            
        Returns:
            bool: True bei Erfolg, False bei Fehler
        """
        try:
            text = ""
            with open(input_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extrahiere Text aus allen Seiten
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"--- Seite {page_num} ---\\n"
                        text += page_text + "\\n\\n"
            
            # Schreibe den extrahierten Text in die Ausgabedatei
            with open(output_path, 'w', encoding='utf-8') as output_file:
                output_file.write(text)
            
            return True
            
        except Exception as e:
            print(f"PDF-Konvertierungsfehler: {e}")
            return False'''

    # converters/office.py
    if folder == 'converters' and filename == 'office.py':
        return '''import os
import subprocess
import tempfile

class OfficeConverter:
    """Konverter für Office-Dokumente"""
    
    def convert(self, input_path, output_path):
        """
        Konvertiert Office-Dokumente zu TXT
        
        Args:
            input_path (str): Pfad zur Eingabedatei
            output_path (str): Pfad zur Ausgabedatei
            
        Returns:
            bool: True bei Erfolg, False bei Fehler
        """
        try:
            # Versuche zuerst mit pandoc
            if self._check_pandoc():
                return self._convert_with_pandoc(input_path, output_path)
            else:
                # Fallback mit Python-Bibliotheken
                return self._fallback_convert(input_path, output_path)
                
        except Exception as e:
            print(f"Office-Konvertierungsfehler: {e}")
            return False
    
    def _check_pandoc(self):
        """Prüft ob pandoc installiert ist"""
        try:
            subprocess.run(['pandoc', '--version'], 
                         capture_output=True, 
                         check=True)
            return True
        except:
            return False
    
    def _convert_with_pandoc(self, input_path, output_path):
        """Konvertiert mit pandoc"""
        try:
            cmd = ['pandoc', input_path, '-o', output_path, '--wrap=none']
            result = subprocess.run(cmd, capture_output=True, text=True)
            
            if result.returncode == 0:
                return True
            else:
                print(f"Pandoc Fehler: {result.stderr}")
                return False
                
        except Exception as e:
            print(f"Pandoc Konvertierungsfehler: {e}")
            return False
    
    def _fallback_convert(self, input_path, output_path):
        """Fallback-Konvertierung mit Python-Bibliotheken"""
        ext = os.path.splitext(input_path)[1].lower()
        
        try:
            if ext == '.docx':
                return self._convert_docx(input_path, output_path)
            elif ext == '.xlsx':
                return self._convert_xlsx(input_path, output_path)
            elif ext == '.odt':
                return self._convert_odt(input_path, output_path)
            elif ext == '.epub':
                return self._convert_epub(input_path, output_path)
            else:
                return False
                
        except ImportError as e:
            print(f"Fallback-Konvertierungsfehler: {e}")
            return False
    
    def _convert_docx(self, input_path, output_path):
        """Konvertiert DOCX zu TXT"""
        try:
            import docx
            doc = docx.Document(input_path)
            
            text = []
            # Extrahiere Text aus Absätzen
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extrahiere Text aus Tabellen
            for table in doc.tables:
                text.append("\\n--- Tabelle ---")
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text.append(" | ".join(row_text))
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\\n'.join(text))
            return True
            
        except Exception as e:
            print(f"DOCX Konvertierungsfehler: {e}")
            return False
    
    def _convert_xlsx(self, input_path, output_path):
        """Konvertiert XLSX zu TXT"""
        try:
            import openpyxl
            wb = openpyxl.load_workbook(input_path, data_only=True)
            text = []
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text.append(f"=== {sheet_name} ===")
                
                for row in sheet.iter_rows(values=True):
                    row_text = []
                    for cell in row:
                        if cell is not None:
                            row_text.append(str(cell))
                    if row_text:
                        text.append('\\t'.join(row_text))
                text.append("")
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\\n'.join(text))
            return True
            
        except Exception as e:
            print(f"XLSX Konvertierungsfehler: {e}")
            return False
    
    def _convert_odt(self, input_path, output_path):
        """Konvertiert ODT zu TXT"""
        try:
            # Versuche mit pandoc, da odt mit Python schwer zu lesen ist
            if self._check_pandoc():
                return self._convert_with_pandoc(input_path, output_path)
            else:
                # Fallback: Versuche mit zipfile den Inhalt zu extrahieren
                import zipfile
                with zipfile.ZipFile(input_path, 'r') as zip_ref:
                    content = zip_ref.read('content.xml').decode('utf-8')
                    # Einfache XML-Text-Extraktion
                    import re
                    text = re.sub(r'<[^>]+>', ' ', content)
                    text = '\\n'.join(line.strip() for line in text.split('\\n') if line.strip())
                    
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(text)
                    return True
        except Exception as e:
            print(f"ODT Konvertierungsfehler: {e}")
            return False
    
    def _convert_epub(self, input_path, output_path):
        """Konvertiert EPUB zu TXT"""
        try:
            # Versuche mit pandoc
            if self._check_pandoc():
                return self._convert_with_pandoc(input_path, output_path)
            else:
                # Fallback: Extrahiere Text aus EPUB
                import zipfile
                with zipfile.ZipFile(input_path, 'r') as zip_ref:
                    text = []
                    for file_info in zip_ref.filelist:
                        if file_info.filename.endswith('.html') or file_info.filename.endswith('.xhtml'):
                            content = zip_ref.read(file_info.filename).decode('utf-8')
                            # Einfache HTML-Text-Extraktion
                            import re
                            clean_text = re.sub(r'<[^>]+>', ' ', content)
                            clean_text = '\\n'.join(line.strip() for line in clean_text.split('\\n') if line.strip())
                            if clean_text:
                                text.append(clean_text)
                    
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write('\\n\\n'.join(text))
                    return True
        except Exception as e:
            print(f"EPUB Konvertierungsfehler: {e}")
            return False'''

    # converters/image.py
    if folder == 'converters' and filename == 'image.py':
        return '''import os
from PIL import Image
import numpy as np

class ImageConverter:
    """Konverter für Bilder mit OCR"""
    
    def __init__(self):
        """Initialisiert den OCR-Reader"""
        try:
            import easyocr
            self.reader = easyocr.Reader(['de', 'en'])
            self.has_ocr = True
        except ImportError:
            print("EasyOCR nicht installiert. OCR wird deaktiviert.")
            self.has_ocr = False
    
    def convert(self, input_path, output_path):
        """
        Konvertiert ein Bild zu TXT (OCR)
        
        Args:
            input_path (str): Pfad zum Bild
            output_path (str): Pfad zur Ausgabedatei
            
        Returns:
            bool: True bei Erfolg, False bei Fehler
        """
        try:
            if not self.has_ocr:
                # Fallback: Versuche mit Tesseract
                return self._convert_with_tesseract(input_path, output_path)
            
            # Lade Bild mit PIL
            image = Image.open(input_path)
            
            # OCR mit EasyOCR
            result = self.reader.readtext(np.array(image), detail=0)
            
            # Schreibe Text
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\\n'.join(result))
            
            return True
            
        except Exception as e:
            print(f"Bild-Konvertierungsfehler: {e}")
            return False
    
    def _convert_with_tesseract(self, input_path, output_path):
        """Fallback mit Tesseract OCR"""
        try:
            import subprocess
            result = subprocess.run(['tesseract', input_path, output_path.replace('.txt', '')], 
                                  capture_output=True, text=True)
            
            if result.returncode == 0 and os.path.exists(output_path):
                return True
            else:
                return False
        except:
            # Wenn Tesseract nicht verfügbar ist, gib einen Fehler zurück
            return False'''

    # converters/text.py
    if folder == 'converters' and filename == 'text.py':
        return '''import os
import json
import xml.etree.ElementTree as ET
import yaml
import csv
import html.parser

class TextConverter:
    """Konverter für Textdateien"""
    
    def convert(self, input_path, output_path):
        """
        Konvertiert verschiedene Textformate zu TXT
        
        Args:
            input_path (str): Pfad zur Eingabedatei
            output_path (str): Pfad zur Ausgabedatei
            
        Returns:
            bool: True bei Erfolg, False bei Fehler
        """
        try:
            ext = os.path.splitext(input_path)[1].lower()
            
            # Lese Datei
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Verarbeite je nach Format
            if ext == '.json':
                text = self._format_json(content)
            elif ext == '.xml':
                text = self._format_xml(content)
            elif ext in ['.yaml', '.yml']:
                text = self._format_yaml(content)
            elif ext == '.csv':
                text = self._format_csv(input_path)
            elif ext == '.html':
                text = self._format_html(content)
            elif ext in ['.py', '.js', '.css', '.bat', '.sh']:
                text = self._format_code(content)
            elif ext in ['.md']:
                text = self._format_markdown(content)
            else:
                text = content
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
            return True
            
        except Exception as e:
            print(f"Text-Konvertierungsfehler: {e}")
            return False
    
    def _format_json(self, content):
        try:
            data = json.loads(content)
            return json.dumps(data, indent=2, ensure_ascii=False)
        except:
            return content
    
    def _format_xml(self, content):
        try:
            root = ET.fromstring(content)
            return ET.tostring(root, encoding='unicode', method='xml')
        except:
            return content
    
    def _format_yaml(self, content):
        try:
            data = yaml.safe_load(content)
            return yaml.dump(data, allow_unicode=True)
        except:
            return content
    
    def _format_csv(self, input_path):
        try:
            text = ""
            with open(input_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                for row in reader:
                    text += '\\t'.join(row) + '\\n'
            return text
        except:
            with open(input_path, 'r', encoding='utf-8') as f:
                return f.read()
    
    def _format_html(self, content):
        # Entferne HTML-Tags
        class HTMLParser(html.parser.HTMLParser):
            def __init__(self):
                super().__init__()
                self.text = []
            
            def handle_data(self, data):
                if data.strip():
                    self.text.append(data.strip())
        
        parser = HTMLParser()
        parser.feed(content)
        return '\\n'.join(parser.text)
    
    def _format_code(self, content):
        # Behalte Code-Formatierung bei
        return content
    
    def _format_markdown(self, content):
        # Behalte Markdown-Formatierung bei
        return content'''

    # converters/archive.py
    if folder == 'converters' and filename == 'archive.py':
        return '''import os
import zipfile

class ArchiveConverter:
    """Konverter für Archive (wird hauptsächlich in app.py verarbeitet)"""
    
    def convert(self, input_path, output_path):
        """
        Archive werden in app.py mit process_zip_file verarbeitet.
        Diese Methode wird normalerweise nicht direkt aufgerufen.
        """
        return False
    
    def extract_archive(self, archive_path, extract_dir):
        """
        Extrahiert ein Archiv in ein Verzeichnis
        
        Args:
            archive_path (str): Pfad zum Archiv
            extract_dir (str): Zielverzeichnis
            
        Returns:
            bool: True bei Erfolg, False bei Fehler
        """
        try:
            with zipfile.ZipFile(archive_path, 'r') as zip_ref:
                zip_ref.extractall(extract_dir)
            return True
        except Exception as e:
            print(f"Archiv-Extraktionsfehler: {e}")
            return False
    
    def get_file_list(self, archive_path):
        """
        Gibt die Liste der Dateien im Archiv zurück
        
        Args:
            archive_path (str): Pfad zum Archiv
            
        Returns:
            list: Liste der Dateinamen
        """
        try:
            with zipfile.ZipFile(archive_path, 'r') as zip_ref:
                return zip_ref.namelist()
        except Exception as e:
            print(f"Fehler beim Lesen des Archivs: {e}")
            return []'''

    # converters/router.py
    if folder == 'converters' and filename == 'router.py':
        return '''import os
from converters.pdf import PDFConverter
from converters.office import OfficeConverter
from converters.image import ImageConverter
from converters.text import TextConverter
from converters.archive import ArchiveConverter

class ConverterRouter:
    def __init__(self):
        """Initialisiert den Router mit allen verfügbaren Konvertern"""
        self.converters = {
            # PDF-Dateien
            '.pdf': PDFConverter(),
            
            # Office-Dokumente
            '.docx': OfficeConverter(),
            '.xlsx': OfficeConverter(),
            '.odt': OfficeConverter(),
            '.epub': OfficeConverter(),
            
            # Bilder (OCR)
            '.jpg': ImageConverter(),
            '.jpeg': ImageConverter(),
            '.png': ImageConverter(),
            '.gif': ImageConverter(),
            '.bmp': ImageConverter(),
            '.tiff': ImageConverter(),
            '.webp': ImageConverter(),
            
            # Textdateien
            '.txt': TextConverter(),
            '.md': TextConverter(),
            '.py': TextConverter(),
            '.json': TextConverter(),
            '.xml': TextConverter(),
            '.yaml': TextConverter(),
            '.yml': TextConverter(),
            '.toml': TextConverter(),
            '.ini': TextConverter(),
            '.cfg': TextConverter(),
            '.csv': TextConverter(),
            '.html': TextConverter(),
            '.css': TextConverter(),
            '.js': TextConverter(),
            '.bat': TextConverter(),
            '.sh': TextConverter(),
            
            # Archive (werden in app.py verarbeitet)
            '.zip': ArchiveConverter(),
        }
    
    def can_convert(self, filename):
        """
        Prüft ob eine Datei konvertiert werden kann
        
        Args:
            filename (str): Name der Datei
            
        Returns:
            bool: True wenn konvertierbar, sonst False
        """
        ext = os.path.splitext(filename)[1].lower()
        return ext in self.converters
    
    def get_converter(self, filename):
        """
        Gibt den passenden Converter für eine Datei zurück
        
        Args:
            filename (str): Name der Datei
            
        Returns:
            object: Converter-Instanz oder None
        """
        ext = os.path.splitext(filename)[1].lower()
        return self.converters.get(ext)
    
    def convert(self, input_path, output_path):
        """
        Konvertiert eine Datei mit dem passenden Converter
        
        Args:
            input_path (str): Pfad zur Eingabedatei
            output_path (str): Pfad zur Ausgabedatei
            
        Returns:
            bool: True bei Erfolg, sonst False
        """
        ext = os.path.splitext(input_path)[1].lower()
        converter = self.converters.get(ext)
        
        if converter:
            try:
                # Spezielle Behandlung für Archive
                if ext == '.zip':
                    # Archive werden in app.py verarbeitet
                    return False
                return converter.convert(input_path, output_path)
            except Exception as e:
                print(f"Konvertierungsfehler für {input_path}: {e}")
                return False
        
        print(f"Kein Converter für {input_path} gefunden")
        return False
    
    def get_supported_extensions(self):
        """
        Gibt alle unterstützten Dateiendungen zurück
        
        Returns:
            list: Liste aller unterstützten Endungen
        """
        return list(self.converters.keys())
    
    def get_supported_formats(self):
        """
        Gibt eine übersichtliche Liste der unterstützten Formate zurück
        
        Returns:
            dict: Dictionary mit Formatgruppen
        """
        formats = {
            'PDF': ['.pdf'],
            'Office': ['.docx', '.xlsx', '.odt', '.epub'],
            'Bilder': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp'],
            'Text': ['.txt', '.md', '.py', '.json', '.xml', '.yaml', '.yml', 
                    '.toml', '.ini', '.cfg', '.csv', '.html', '.css', '.js', 
                    '.bat', '.sh'],
            'Archive': ['.zip']
        }
        return formats
    
    def is_archive(self, filename):
        """
        Prüft ob es sich um eine Archiv-Datei handelt
        
        Args:
            filename (str): Name der Datei
            
        Returns:
            bool: True wenn Archiv, sonst False
        """
        ext = os.path.splitext(filename)[1].lower()
        return ext == '.zip'''

    # templates/index.html
    if folder == 'templates' and filename == 'index.html':
        return '''<!DOCTYPE html>
<html lang="de">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>AI File Platform</title>
    <style>
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }
        
        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            min-height: 100vh;
            padding: 20px;
        }
        
        .container {
            max-width: 1200px;
            margin: 0 auto;
            background: white;
            border-radius: 20px;
            box-shadow: 0 20px 60px rgba(0,0,0,0.3);
            padding: 40px;
        }
        
        h1 {
            color: #333;
            font-size: 2.5em;
            margin-bottom: 10px;
            text-align: center;
        }
        
        .subtitle {
            text-align: center;
            color: #666;
            margin-bottom: 30px;
            font-size: 1.1em;
        }
        
        .upload-area {
            border: 3px dashed #ddd;
            border-radius: 15px;
            padding: 60px 20px;
            text-align: center;
            transition: all 0.3s ease;
            background: #fafafa;
            cursor: pointer;
            position: relative;
        }
        
        .upload-area:hover, .upload-area.dragover {
            border-color: #667eea;
            background: #f0f0ff;
            transform: scale(1.01);
        }
        
        .upload-icon {
            font-size: 4em;
            margin-bottom: 20px;
        }
        
        .upload-text {
            font-size: 1.2em;
            color: #555;
        }
        
        .upload-subtext {
            color: #999;
            margin-top: 10px;
        }
        
        #fileInput {
            display: none;
        }
        
        .file-list {
            margin-top: 30px;
        }
        
        .file-item {
            background: #f8f9fa;
            border-radius: 8px;
            padding: 15px 20px;
            margin-bottom: 10px;
            display: flex;
            justify-content: space-between;
            align-items: center;
            animation: slideIn 0.3s ease;
        }
        
        .file-name {
            display: flex;
            align-items: center;
            gap: 10px;
        }
        
        .file-status {
            padding: 5px 12px;
            border-radius: 20px;
            font-size: 0.85em;
            font-weight: 500;
        }
        
        .status-pending {
            background: #ffd93d;
            color: #8b6f00;
        }
        
        .status-processing {
            background: #6c5ce7;
            color: white;
        }
        
        .status-success {
            background: #00b894;
            color: white;
        }
        
        .status-error {
            background: #ff6b6b;
            color: white;
        }
        
        .status-skipped {
            background: #dfe6e9;
            color: #2d3436;
        }
        
        .progress-container {
            margin-top: 30px;
            background: #f0f0f0;
            border-radius: 10px;
            height: 30px;
            overflow: hidden;
            position: relative;
            display: none;
        }
        
        .progress-bar {
            height: 100%;
            background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
            transition: width 0.5s ease;
            width: 0%;
            border-radius: 10px;
        }
        
        .progress-text {
            position: absolute;
            top: 50%;
            left: 50%;
            transform: translate(-50%, -50%);
            color: #333;
            font-weight: 600;
            font-size: 0.9em;
        }
        
        .download-section {
            margin-top: 30px;
            text-align: center;
            display: none;
        }
        
        .btn {
            display: inline-block;
            padding: 12px 30px;
            border: none;
            border-radius: 8px;
            font-size: 1em;
            font-weight: 600;
            cursor: pointer;
            transition: all 0.3s ease;
        }
        
        .btn-primary {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
        }
        
        .btn-primary:hover {
            transform: translateY(-2px);
            box-shadow: 0 5px 20px rgba(102, 126, 234, 0.4);
        }
        
        .btn-success {
            background: #00b894;
            color: white;
        }
        
        .btn-success:hover {
            transform: translateY(-2px);
            box-shadow: 0 5px 20px rgba(0, 184, 148, 0.4);
        }
        
        .error-list {
            margin-top: 20px;
            background: #fff5f5;
            border-left: 4px solid #ff6b6b;
            padding: 15px;
            border-radius: 5px;
            display: none;
        }
        
        .error-item {
            color: #c0392b;
            font-size: 0.9em;
            padding: 5px 0;
        }
        
        .preview-section {
            margin-top: 30px;
            background: #f8f9fa;
            border-radius: 8px;
            padding: 20px;
            display: none;
            max-height: 400px;
            overflow-y: auto;
        }
        
        .preview-content {
            white-space: pre-wrap;
            font-family: 'Courier New', monospace;
            font-size: 0.9em;
        }
        
        @keyframes slideIn {
            from {
                opacity: 0;
                transform: translateY(-10px);
            }
            to {
                opacity: 1;
                transform: translateY(0);
            }
        }
        
        .supported-formats {
            margin-top: 30px;
            padding: 20px;
            background: #f8f9fa;
            border-radius: 10px;
        }
        
        .format-tag {
            display: inline-block;
            background: #e9ecef;
            padding: 3px 10px;
            border-radius: 15px;
            font-size: 0.85em;
            margin: 3px;
            color: #495057;
        }
        
        @media (max-width: 768px) {
            .container {
                padding: 20px;
            }
            
            h1 {
                font-size: 1.8em;
            }
            
            .upload-area {
                padding: 30px 15px;
            }
        }
    </style>
</head>
<body>
    <div class="container">
        <h1>📄 AI File Platform</h1>
        <p class="subtitle">Konvertiere verschiedene Dateiformate in TXT</p>
        
        <div class="upload-area" id="dropZone">
            <div class="upload-icon">📁</div>
            <div class="upload-text">Dateien hier ablegen oder klicken zum Hochladen</div>
            <div class="upload-subtext">Unterstützt alle gängigen Formate • Max 500MB</div>
            <input type="file" id="fileInput" multiple>
        </div>
        
        <div class="progress-container" id="progressContainer">
            <div class="progress-bar" id="progressBar"></div>
            <div class="progress-text" id="progressText">0%</div>
        </div>
        
        <div class="file-list" id="fileList"></div>
        
        <div class="error-list" id="errorList"></div>
        
        <div class="preview-section" id="previewSection">
            <h4>📝 Vorschau</h4>
            <div class="preview-content" id="previewContent"></div>
        </div>
        
        <div class="download-section" id="downloadSection">
            <button class="btn btn-success" id="downloadAllBtn">⬇️ Alles als TXT herunterladen</button>
        </div>
        
        <div class="supported-formats">
            <strong>Unterstützte Formate:</strong>
            <span class="format-tag">.txt</span>
            <span class="format-tag">.md</span>
            <span class="format-tag">.py</span>
            <span class="format-tag">.json</span>
            <span class="format-tag">.xml</span>
            <span class="format-tag">.yaml</span>
            <span class="format-tag">.csv</span>
            <span class="format-tag">.html</span>
            <span class="format-tag">.css</span>
            <span class="format-tag">.js</span>
            <span class="format-tag">.pdf</span>
            <span class="format-tag">.docx</span>
            <span class="format-tag">.xlsx</span>
            <span class="format-tag">.jpg</span>
            <span class="format-tag">.png</span>
            <span class="format-tag">.gif</span>
            <span class="format-tag">.zip</span>
        </div>
    </div>
    
    <script>
        let sessionId = null;
        let statusInterval = null;
        
        const dropZone = document.getElementById('dropZone');
        const fileInput = document.getElementById('fileInput');
        const fileList = document.getElementById('fileList');
        const progressContainer = document.getElementById('progressContainer');
        const progressBar = document.getElementById('progressBar');
        const progressText = document.getElementById('progressText');
        const downloadSection = document.getElementById('downloadSection');
        const downloadAllBtn = document.getElementById('downloadAllBtn');
        const errorList = document.getElementById('errorList');
        const previewSection = document.getElementById('previewSection');
        const previewContent = document.getElementById('previewContent');
        
        // Drag & Drop Events
        dropZone.addEventListener('dragover', (e) => {
            e.preventDefault();
            dropZone.classList.add('dragover');
        });
        
        dropZone.addEventListener('dragleave', () => {
            dropZone.classList.remove('dragover');
        });
        
        dropZone.addEventListener('drop', (e) => {
            e.preventDefault();
            dropZone.classList.remove('dragover');
            const files = e.dataTransfer.files;
            handleFiles(files);
        });
        
        dropZone.addEventListener('click', () => {
            fileInput.click();
        });
        
        fileInput.addEventListener('change', (e) => {
            handleFiles(e.target.files);
            fileInput.value = '';
        });
        
        function handleFiles(files) {
            if (files.length === 0) return;
            
            const formData = new FormData();
            for (let file of files) {
                formData.append('files', file);
            }
            
            // Clear previous results
            fileList.innerHTML = '';
            errorList.style.display = 'none';
            errorList.innerHTML = '';
            downloadSection.style.display = 'none';
            previewSection.style.display = 'none';
            
            // Show progress
            progressContainer.style.display = 'block';
            progressBar.style.width = '0%';
            progressText.textContent = '0%';
            
            // Upload files
            fetch('/upload', {
                method: 'POST',
                body: formData
            })
            .then(response => response.json())
            .then(data => {
                sessionId = data.session_id;
                startStatusPolling();
            })
            .catch(error => {
                console.error('Upload error:', error);
                showError('Fehler beim Hochladen der Dateien');
            });
        }
        
        function startStatusPolling() {
            if (statusInterval) clearInterval(statusInterval);
            
            statusInterval = setInterval(() => {
                fetch(`/status/${sessionId}`)
                    .then(response => response.json())
                    .then(data => {
                        updateUI(data);
                        
                        if (data.status === 'completed' || data.status === 'error') {
                            clearInterval(statusInterval);
                            if (data.status === 'completed' && data.converted_files && data.converted_files.length > 0) {
                                downloadSection.style.display = 'block';
                            }
                        }
                    })
                    .catch(error => {
                        console.error('Status error:', error);
                    });
            }, 1000);
        }
        
        function updateUI(data) {
            // Update progress
            progressBar.style.width = data.progress + '%';
            progressText.textContent = data.progress + '%';
            
            // Update file list
            fileList.innerHTML = '';
            data.files.forEach(file => {
                const fileItem = document.createElement('div');
                fileItem.className = 'file-item';
                
                const fileName = document.createElement('div');
                fileName.className = 'file-name';
                fileName.innerHTML = `<span>${getFileIcon(file.name)}</span> ${file.name}`;
                
                const status = document.createElement('span');
                status.className = `file-status status-${file.status}`;
                const statusText = {
                    'pending': '⏳ Ausstehend',
                    'processing': '🔄 Verarbeitung',
                    'success': '✅ Erfolgreich',
                    'error': '❌ Fehler',
                    'skipped': '⏭️ Übersprungen'
                };
                status.textContent = statusText[file.status] || file.status;
                
                fileItem.appendChild(fileName);
                fileItem.appendChild(status);
                fileList.appendChild(fileItem);
            });
            
            // Update errors
            if (data.errors && data.errors.length > 0) {
                errorList.style.display = 'block';
                errorList.innerHTML = '<h4>⚠️ Fehler:</h4>';
                data.errors.forEach(error => {
                    const errorItem = document.createElement('div');
                    errorItem.className = 'error-item';
                    errorItem.textContent = error;
                    errorList.appendChild(errorItem);
                });
            }
        }
        
        function getFileIcon(filename) {
            const ext = filename.split('.').pop().toLowerCase();
            const icons = {
                'pdf': '📄',
                'docx': '📝',
                'xlsx': '📊',
                'txt': '📃',
                'md': '📝',
                'py': '🐍',
                'json': '📋',
                'xml': '📋',
                'yaml': '📋',
                'yml': '📋',
                'csv': '📊',
                'html': '🌐',
                'css': '🎨',
                'js': '📜',
                'jpg': '🖼️',
                'jpeg': '🖼️',
                'png': '🖼️',
                'gif': '🖼️',
                'zip': '📦'
            };
            return icons[ext] || '📎';
        }
        
        function showError(message) {
            errorList.style.display = 'block';
            const errorItem = document.createElement('div');
            errorItem.className = 'error-item';
            errorItem.textContent = message;
            errorList.appendChild(errorItem);
        }
        
        // Download all
        downloadAllBtn.addEventListener('click', () => {
            if (sessionId) {
                window.location.href = `/download/${sessionId}`;
            }
        });
        
        // Preview functionality
        fileList.addEventListener('click', (e) => {
            const fileItem = e.target.closest('.file-item');
            if (!fileItem) return;
            
            const fileNameElement = fileItem.querySelector('.file-name span');
            if (!fileNameElement) return;
            
            const fileName = fileNameElement.textContent || '';
            if (!fileName) return;
            
            // Find converted file by checking status
            fetch(`/status/${sessionId}`)
                .then(response => response.json())
                .then(data => {
                    // Find the converted file
                    const converted = data.converted_files.find(f => 
                        f.original === fileName
                    );
                    
                    if (converted) {
                        fetch(`/preview/${sessionId}/${converted.converted}`)
                            .then(response => response.json())
                            .then(data => {
                                if (data.content) {
                                    previewSection.style.display = 'block';
                                    previewContent.textContent = data.content;
                                } else if (data.error) {
                                    alert('Vorschau Fehler: ' + data.error);
                                }
                            })
                            .catch(error => {
                                console.error('Preview error:', error);
                            });
                    } else {
                        // Check if file is already converted
                        const statusElement = fileItem.querySelector('.file-status');
                        if (statusElement && statusElement.textContent.includes('Erfolgreich')) {
                            // File is converted but not in list yet
                            setTimeout(() => {
                                fileItem.click();
                            }, 1000);
                        }
                    }
                });
        });
    </script>
</body>
</html>'''

    # requirements.txt
    if filename == 'requirements.txt':
        return '''Flask==3.0.0
Werkzeug==3.0.1
PyPDF2==3.0.1
Pillow==10.1.0
easyocr==1.7.0
python-docx==1.1.0
openpyxl==3.1.2
PyYAML==6.0.1
numpy==1.24.3'''

# README.md
return '''# AI File Platform

Eine Webanwendung zur Konvertierung verschiedener Dateiformate in TXT.

## Funktionen

- ✅ Drag-and-Drop Upload
- ✅ Mehrere Dateien gleichzeitig
- ✅ ZIP-Dateien automatisch entpacken
- ✅ PDF → TXT
- ✅ DOCX → TXT
- ✅ XLSX → TXT
- ✅ HTML → TXT
- ✅ JSON → TXT
- ✅ Bilder → TXT (OCR)
- ✅ Python-Dateien → TXT
- ✅ Markdown-Dateien → TXT
- ✅ Doppelte Dateien erkennen
- ✅ "Alles als TXT herunterladen"
- ✅ Fortschrittsanzeige
- ✅ Dateivorschau

## Installation

```bash
git clone <repository>
cd ai-file-platform
python -m venv venv